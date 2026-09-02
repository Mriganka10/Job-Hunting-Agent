import json
from pathlib import Path

from docx import Document

from job_hunting_agent import ats_semantic
from job_hunting_agent.ats import score_resume
from job_hunting_agent.ats_calibration import calibration_cases, load_calibration_manifest, run_calibration_benchmark
from job_hunting_agent.ats_layout import analyze_resume_layout
from job_hunting_agent.ats_profiles import select_role_profile
from job_hunting_agent.ats_semantic import semantic_similarity
from job_hunting_agent.models import CandidateProfile, Resume
from job_hunting_agent.resume import extract_sections


def test_calibration_corpus_has_sixty_role_diverse_anonymized_cases() -> None:
    cases = calibration_cases()

    assert len(cases) == 60
    assert len({case.role_family for case in cases}) == 10
    assert {case.quality_tier for case in cases} == {"weak", "basic", "developing", "solid", "strong", "excellent"}
    assert all(case.provenance == "anonymized_role_archetype" for case in cases)


def test_calibration_benchmark_meets_v1_coverage_baseline() -> None:
    result = run_calibration_benchmark()

    assert result["case_count"] == 60
    assert result["band_coverage"] >= 0.90
    assert result["mean_absolute_error"] <= 6.0


def test_private_calibration_manifest_loads_resume_without_copying_it_into_repo(tmp_path: Path) -> None:
    resume_path = tmp_path / "reviewed.txt"
    resume_path.write_text("CONTACT\ncandidate@example.com\nSKILLS\nPython\nEXPERIENCE\nBuilt reliable APIs.\nEDUCATION\nB.Tech", encoding="utf-8")
    manifest_path = tmp_path / "calibration.json"
    manifest_path.write_text(
        json.dumps(
            {
                "cases": [
                    {
                        "case_id": "reviewed-001",
                        "resume_path": "reviewed.txt",
                        "role_family": "software_engineering",
                        "target_role": "Software Engineer",
                        "expected_min": 45,
                        "expected_max": 60,
                    }
                ]
            }
        ),
        encoding="utf-8",
    )

    cases = load_calibration_manifest(manifest_path)

    assert len(cases) == 1
    assert cases[0].case_id == "reviewed-001"
    assert cases[0].provenance == "consented_private_manifest"
    assert "candidate@example.com" in cases[0].resume_text


def test_role_profile_selection_uses_explicit_role_and_exposes_signals() -> None:
    profile, confidence, signals = select_role_profile(
        ("Adjunct Faculty",),
        ("Financial Analyst",),
        "Teach corporate finance and mentor postgraduate students.",
    )

    assert profile.key == "academic_teaching"
    assert sum(profile.category_weights) == 100
    assert confidence >= 0.70
    assert any("target role" in signal for signal in signals)


def test_score_exposes_confidence_range_role_and_semantic_provenance(tmp_path: Path) -> None:
    text = """CONTACT
candidate@example.com | +91 9000000000
PROFESSIONAL SUMMARY
Data Engineer building reliable batch and streaming pipelines with Python and SQL.
TECHNICAL SKILLS
Python, SQL, Spark, ETL, AWS
PROFESSIONAL EXPERIENCE
Data Engineer | Example Company | January 2021 - Present
- Built ETL pipelines processing 2,000,000 records and reduced runtime by 32%.
- Improved data quality controls for 18 production datasets.
PROJECTS
- Designed a Spark pipeline serving 300 analysts.
EDUCATION
Bachelor of Technology | 2020
"""
    path = tmp_path / "resume.txt"
    path.write_text(text, encoding="utf-8")
    resume = Resume(str(path), text, ("Python", "SQL", "Spark", "ETL", "AWS"), ("Data Engineer",), extract_sections(text))
    report = score_resume(
        resume,
        CandidateProfile(
            target_roles=("Data Engineer",),
            skills=("Python", "SQL", "Spark", "ETL", "AWS"),
            job_description="Data Engineer building Spark ETL pipelines with Python, SQL, and AWS.",
            experience_years=4,
        ),
    )

    assert report.role_profile == "Data Engineering"
    assert report.confidence_label in {"Moderate", "High"}
    assert report.score_range[0] <= report.score <= report.score_range[1]
    assert report.extraction_confidence >= 0.70
    assert report.semantic_provider == "lexical_tfidf_fallback"
    assert report.layout_analysis["checks"]["status"] == "plain_text"
    assert report.category_details["role_profile"]["key"] == "data_engineering"


def test_semantic_fallback_is_tfidf_not_hashed(monkeypatch) -> None:
    monkeypatch.delenv("JOB_AGENT_EMBEDDING_API_KEY", raising=False)
    monkeypatch.delenv("JOB_AGENT_EMBEDDING_ENDPOINT", raising=False)
    monkeypatch.setenv("JOB_AGENT_ENABLE_LOCAL_EMBEDDINGS", "false")
    semantic_similarity.cache_clear()

    related = semantic_similarity("Spark ETL data pipeline on AWS", "Build AWS data pipelines using Spark ETL")
    unrelated = semantic_similarity("Spark ETL data pipeline on AWS", "Teach medieval art history and sculpture")

    assert related.provider == "lexical_tfidf_fallback"
    assert related.similarity > unrelated.similarity
    assert related.confidence < 0.80


def test_trained_semantic_provider_reports_model_provenance(monkeypatch) -> None:
    class FakeModel:
        def encode(self, values, normalize_embeddings=True):
            assert len(values) == 2
            assert normalize_embeddings is True
            return ([1.0, 0.0, 0.0], [0.8, 0.6, 0.0])

    monkeypatch.setattr(ats_semantic, "_sentence_transformer_model", lambda: FakeModel())
    result = ats_semantic._sentence_transformer_similarity("data engineering", "data pipelines")

    assert result is not None
    assert result.provider == "sentence_transformer"
    assert result.similarity == 0.8
    assert result.confidence >= 0.90


def test_docx_layout_check_detects_ats_hostile_structures(tmp_path: Path) -> None:
    path = tmp_path / "table-resume.docx"
    document = Document()
    document.sections[0].header.paragraphs[0].text = "candidate@example.com"
    document.add_paragraph("PROFESSIONAL SUMMARY")
    document.add_paragraph("Data analyst with SQL and reporting experience.")
    for value in ("Skills", "Experience", "Education"):
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = value
        table.cell(0, 1).text = "Structured content"
    document.save(path)

    result = analyze_resume_layout(path, "Readable extracted resume text " * 30)

    assert result["format"] == "docx"
    assert result["score"] < 100
    assert any("tables" in issue.casefold() for issue in result["issues"])
    assert any("header" in issue.casefold() for issue in result["issues"])
