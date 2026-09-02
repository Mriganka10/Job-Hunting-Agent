from __future__ import annotations

import re
import unicodedata
from pathlib import Path

from .models import Resume

KNOWN_SKILLS = {
    "python",
    "java",
    "javascript",
    "typescript",
    "sql",
    "pl/sql",
    "fastapi",
    "django",
    "flask",
    "react",
    "react.js",
    "node",
    "node.js",
    "aws",
    "azure",
    "gcp",
    "docker",
    "kubernetes",
    "machine learning",
    "deep learning",
    "nlp",
    "pandas",
    "spark",
    "pyspark",
    "airflow",
    "terraform",
    "linux",
    "unix",
    "bash",
    "c++", "mysql", "postgresql", "mongodb", "oracle", "power bi", "tableau",
    "tensorflow", "keras", "pytorch", "scikit-learn",
    "numpy", "matplotlib", "seaborn", "opencv", "nltk", "spacy", "github", "gitlab", "git",
    "github actions",
    "jupyter notebook", "data structures", "algorithms", "dbms",
    "computer networks", "object detection", "image classification",
    "rest api", "microservices", "ci/cd", "jenkins", "jira", "autosys",
    "hadoop", "hive", "sqoop", "kafka", "databricks", "dbt", "trino",
    "delta lake", "apache nifi", "excel",
    "artificial intelligence", "ai", "ml", "generative ai", "agentic ai",
    "rpa", "uipath", "vba", "rnn", "lstm", "eda", "dvc", "oop",
    "operating systems", "vs code", "pycharm", "named entity recognition",
    "text classification", "financial modelling", "market sizing",
    "business economics", "corporate strategy", "m&a research", "m&a",
    "investment banking research", "secondary research", "market research",
    "stock market fundamental analysis", "technical analysis",
    "competitive intelligence", "financial analysis", "business research",
    "data analytics", "bloomberg", "cap iq", "factset", "thomson eikon",
    "eikon", "dealogic", "sdc", "euromonitor", "gartner", "idc", "eiu",
    "moody's", "fitch", "factiva", "ims", "oxford economics", "finbert",
    "dspy", "scala", "hadoop mapreduce", "pig", "svn", "windows",
    "solaris", "azure data factory", "azure datalake gen2",
    "azure synapse analytics", "unix shell script", "oracle 11g",
    "ms excel",
}

CANONICAL_SKILL_LABELS = {
    "airflow": "Airflow",
    "algorithms": "Algorithms",
    "apache nifi": "Apache NiFi",
    "agentic ai": "Agentic AI",
    "ai": "AI",
    "artificial intelligence": "Artificial Intelligence",
    "autosys": "AutoSys",
    "aws": "AWS",
    "azure": "Azure",
    "bash": "Bash",
    "bloomberg": "Bloomberg",
    "business economics": "Business Economics",
    "business research": "Business Research",
    "c plus plus": "C++",
    "cap iq": "Cap IQ",
    "ci cd": "CI/CD",
    "competitive intelligence": "Competitive Intelligence",
    "computer networks": "Computer Networks",
    "corporate strategy": "Corporate Strategy",
    "data analytics": "Data Analytics",
    "data structures": "Data Structures",
    "databricks": "Databricks",
    "dbms": "DBMS",
    "dbt": "dbt",
    "deep learning": "Deep Learning",
    "dealogic": "Dealogic",
    "delta lake": "Delta Lake",
    "django": "Django",
    "docker": "Docker",
    "dspy": "DSPy",
    "dvc": "DVC",
    "eda": "EDA",
    "eikon": "Eikon",
    "eiu": "EIU",
    "excel": "Excel",
    "euromonitor": "Euromonitor",
    "factiva": "Factiva",
    "factset": "FactSet",
    "fastapi": "FastAPI",
    "financial analysis": "Financial Analysis",
    "financial modelling": "Financial Modelling",
    "finbert": "FinBERT",
    "fitch": "Fitch",
    "flask": "Flask",
    "gartner": "Gartner",
    "gcp": "GCP",
    "generative ai": "Generative AI",
    "git": "Git",
    "github": "GitHub",
    "github actions": "GitHub Actions",
    "gitlab": "GitLab",
    "hadoop": "Hadoop",
    "hadoop mapreduce": "Hadoop MapReduce",
    "hive": "Hive",
    "idc": "IDC",
    "image classification": "Image Classification",
    "ims": "IMS",
    "investment banking research": "Investment Banking Research",
    "java": "Java",
    "javascript": "JavaScript",
    "jenkins": "Jenkins",
    "jira": "Jira",
    "jupyter notebook": "Jupyter Notebook",
    "keras": "Keras",
    "kafka": "Kafka",
    "kubernetes": "Kubernetes",
    "linux": "Linux",
    "lstm": "LSTM",
    "m a": "M&A",
    "m&a": "M&A",
    "m a research": "M&A Research",
    "m&a research": "M&A Research",
    "machine learning": "Machine Learning",
    "market research": "Market Research",
    "market sizing": "Market Sizing",
    "matplotlib": "Matplotlib",
    "microservices": "Microservices",
    "ml": "ML",
    "mongodb": "MongoDB",
    "ms excel": "MS Excel",
    "mysql": "MySQL",
    "named entity recognition": "Named Entity Recognition",
    "nlp": "NLP",
    "node": "Node.js",
    "node js": "Node.js",
    "nltk": "NLTK",
    "numpy": "NumPy",
    "object detection": "Object Detection",
    "oop": "OOP",
    "opencv": "OpenCV",
    "operating systems": "Operating Systems",
    "oracle": "Oracle",
    "oracle 11g": "Oracle 11g",
    "oxford economics": "Oxford Economics",
    "pandas": "Pandas",
    "pl sql": "PL/SQL",
    "postgresql": "PostgreSQL",
    "power bi": "Power BI",
    "pig": "PIG",
    "pyspark": "PySpark",
    "pycharm": "PyCharm",
    "python": "Python",
    "pytorch": "PyTorch",
    "react": "React",
    "react js": "React.js",
    "rest api": "REST API",
    "rnn": "RNN",
    "rpa": "RPA",
    "scikit learn": "scikit-learn",
    "scala": "Scala",
    "seaborn": "Seaborn",
    "secondary research": "Secondary Research",
    "sdc": "SDC",
    "spark": "Spark",
    "spacy": "spaCy",
    "sqoop": "Sqoop",
    "sql": "SQL",
    "solaris": "Solaris",
    "svn": "SVN",
    "tableau": "Tableau",
    "technical analysis": "Technical Analysis",
    "tensorflow": "TensorFlow",
    "terraform": "Terraform",
    "text classification": "Text Classification",
    "thomson eikon": "Thomson Eikon",
    "trino": "Trino",
    "typescript": "TypeScript",
    "unix": "Unix",
    "unix shell script": "Unix Shell Script",
    "uipath": "UiPath",
    "vba": "VBA",
    "vs code": "VS Code",
    "windows": "Windows",
    "azure data factory": "Azure Data Factory",
    "azure datalake gen2": "Azure DataLake Gen2",
    "azure synapse analytics": "Azure Synapse Analytics",
}

GENERIC_SKILL_LABELS = {
    "additional skills",
    "analytics",
    "ai and technology",
    "ai technology",
    "big data technology",
    "cloud",
    "cloud and devops",
    "cloud technology",
    "competencies",
    "core skills",
    "data platforms",
    "data visualization tool",
    "database",
    "databases",
    "development tools",
    "devops tool",
    "domain expertise",
    "eagle technology",
    "frameworks",
    "languages",
    "methodologies",
    "operating systems",
    "platforms",
    "programming",
    "programming language",
    "research database",
    "research databases",
    "scripting language",
    "skills",
    "source versioning control tool",
    "technical competencies",
    "technical skills",
    "technologies",
    "tools",
    "tools and technologies",
}

ROLE_HINTS = (
    "ai developer",
    "data scientist",
    "data analyst",
    "ml engineer",
    "software engineer",
    "python developer",
    "machine learning engineer",
    "data engineer",
    "backend developer",
    "full stack developer",
    "devops engineer",
)

SECTION_ALIASES = {
    "contact": ("contact", "contact details", "personal details"),
    "summary": ("summary", "professional summary", "career summary", "profile", "profile summary", "objective", "career objective", "carrer objective", "about me"),
    "experience": ("experience", "work experience", "professional experience", "corporate experience", "employment", "employment history", "work history", "professional history", "internship", "internships"),
    "skills": ("skills", "technical skills", "core skills", "core skills and tools", "key skills", "competencies", "technical competencies", "technologies", "tools and technologies", "domain expertise"),
    "education": ("education", "academic credentials", "academic background", "academic qualifications", "educational qualifications", "qualifications"),
    "projects": ("projects", "project experience", "academic projects", "personal projects", "key projects", "portfolio"),
    "certifications": ("certifications", "certification", "licenses and certifications", "courses", "training"),
    "achievements": ("achievements", "achievement", "key achievements", "corporate achievements", "key corporate achievements", "awards", "honors", "accomplishments", "achievement and certification", "achievements and certifications", "achievement certification", "achievements certifications"),
    "languages": ("languages", "language", "languages known", "language proficiency"),
    "publications": ("publications", "publication", "research publications", "research papers"),
    "volunteering": ("volunteering", "volunteer experience", "community involvement", "leadership experience"),
    "interests": ("interests", "areas of interest", "professional interests", "hobbies"),
    "core_competencies": ("core competencies", "areas of expertise", "key competencies"),
    "soft_skills": ("soft skills", "professional strengths", "personal skills", "interpersonal skills"),
    "career_timeline": ("career timeline", "employment timeline", "career progression"),
    "teaching_vision": ("teaching vision", "teaching philosophy", "statement of teaching philosophy"),
    "teaching_subjects": ("subjects available to teach", "subjects taught", "courses available to teach", "courses taught", "teaching areas"),
}


def parse_resume(path: str | Path) -> Resume:
    resume_path = Path(path)
    if not resume_path.exists():
        raise FileNotFoundError(f"Resume not found: {resume_path}")

    suffix = resume_path.suffix.lower()
    if suffix == ".pdf":
        text = _read_pdf(resume_path)
    elif suffix == ".docx":
        text = _read_docx(resume_path)
    elif suffix in {".txt", ".md"}:
        text = resume_path.read_text(encoding="utf-8")
    else:
        raise ValueError("Supported resume formats are PDF, DOCX, TXT, and MD.")

    normalized = normalize_text(text)
    return Resume(
        path=str(resume_path),
        text=normalized,
        inferred_skills=extract_skills(normalized),
        inferred_roles=extract_roles(normalized),
        sections=extract_sections(normalized),
    )


def extract_sections(text: str) -> dict[str, str]:
    """Split normalized resume text into canonical ATS sections."""
    lines = text.splitlines()
    aliases = {
        _heading_key(alias): section
        for section, names in SECTION_ALIASES.items()
        for alias in names
    }
    sections: dict[str, list[str]] = {name: [] for name in SECTION_ALIASES}
    current = "contact"
    for raw_line in lines:
        cleaned = re.sub(r"^[\s\-–—•▪●*#|]+", "", raw_line).strip()
        heading_part = re.split(r"\s*[:|]\s*", cleaned, maxsplit=1)[0]
        key = _heading_key(heading_part)
        matched = aliases.get(key)
        if not matched and len(cleaned) <= 80:
            matched = next((section for alias, section in aliases.items() if key.startswith(f"{alias} ")), None)
        if matched:
            if matched == current and cleaned != cleaned.upper():
                sections[current].append(raw_line)
                continue
            current = matched
            remainder = re.sub(rf"^{re.escape(heading_part)}\s*[:|\-–—]?\s*", "", cleaned, flags=re.I).strip()
            if remainder:
                sections[current].append(remainder)
            continue
        sections[current].append(raw_line)
    result = {name: "\n".join(items).strip() for name, items in sections.items() if "\n".join(items).strip()}
    result["contact"] = _contact_block(lines)
    if not result.get("education"):
        recovered = _recover_education(lines)
        if recovered:
            result["education"] = recovered
    if result.get("achievements") and not result.get("certifications"):
        certification_lines = [
            line for line in result["achievements"].splitlines()
            if re.search(r"\b(?:certif|training|course|license)\w*\b", line, flags=re.I)
        ]
        if certification_lines:
            result["certifications"] = "\n".join(certification_lines)
    if result.get("certifications") and not result.get("achievements"):
        achievement_lines = [
            line for line in result["certifications"].splitlines()
            if re.search(
                r"\b(?:award|winner|won|topper|honou?r|recogniz|achievement|accomplish|"
                r"rank(?:ed)?|solved|increased|reduced|improved)\b|\d+\s*[+%]",
                line,
                flags=re.I,
            )
            and not re.search(r"\b(?:certif|training|course|license)\w*\b", line, flags=re.I)
        ]
        if achievement_lines:
            result["achievements"] = "\n".join(achievement_lines)
    return result


def normalize_text(text: str) -> str:
    text = unicodedata.normalize("NFKC", text).replace("\u00ad", "").replace("�", "-")
    text = re.sub(r"(?<=\w)-\s*\n\s*(?=\w)", "", text)
    lines = [normalize_ats_text(_repair_spaced_glyphs(line)) for line in text.splitlines()]
    normalized = "\n".join(line for line in lines if line)
    return re.sub(r"\n{3,}", "\n\n", normalized).strip()


def normalize_ats_text(value: str) -> str:
    """Normalize common ATS keyword casing and resume extraction glitches."""
    replacements = {
        "highquality": "high-quality",
        "high quality": "high-quality",
        "powerbi": "Power BI",
    }
    normalized = value
    for source, target in sorted(replacements.items(), key=lambda item: len(item[0]), reverse=True):
        normalized = re.sub(
            rf"(?<![A-Za-z0-9]){re.escape(source)}(?![A-Za-z0-9])",
            target,
            normalized,
            flags=re.I,
        )
    for key, label in sorted(CANONICAL_SKILL_LABELS.items(), key=lambda item: len(item[0]), reverse=True):
        if key in {"c plus plus", "ci cd", "pl sql", "react js", "node js", "scikit learn"}:
            continue
        normalized = re.sub(
            rf"(?<![A-Za-z0-9+#.-]){re.escape(key)}(?![A-Za-z0-9+#.-])",
            label,
            normalized,
            flags=re.I,
        )
    normalized = re.sub(r"\bci\s*/?\s*cd\b", "CI/CD", normalized, flags=re.I)
    normalized = re.sub(r"\bpl\s*/\s*sql\b", "PL/SQL", normalized, flags=re.I)
    normalized = re.sub(r"\bnode\s*\.?\s*js\b", "Node.js", normalized, flags=re.I)
    normalized = re.sub(r"\breact\s*\.?\s*js\b", "React.js", normalized, flags=re.I)
    normalized = re.sub(r"\bscikit\s+learn\b", "scikit-learn", normalized, flags=re.I)
    return normalized


def _repair_spaced_glyphs(line: str) -> str:
    """Repair PDFs that encode every visible character as a separate word."""
    raw = line.replace("\u00a0", " ").strip()
    tokens = raw.split()
    if len(tokens) < 4:
        return re.sub(r"[ \t]+", " ", raw)
    glyph_tokens = sum(1 for token in tokens if len(re.sub(r"[^A-Za-z0-9]", "", token)) <= 1)
    if glyph_tokens / len(tokens) < 0.65:
        return re.sub(r"[ \t]+", " ", raw)
    words = []
    for group in re.split(r"[ \t]{2,}", raw):
        if group.strip():
            words.append(re.sub(r"(?<=\S)[ \t](?=\S)", "", group.strip()))
    return " ".join(words)


def _contact_block(lines: list[str]) -> str:
    heading_keys = {_heading_key(alias) for names in SECTION_ALIASES.values() for alias in names}
    name_candidates: list[str] = []
    contact_values: list[str] = []
    for index, line in enumerate(lines):
        cleaned = line.strip()
        if not cleaned or _heading_key(cleaned) in heading_keys:
            continue
        if index < 16 and _looks_like_name(cleaned):
            name_candidates.append(cleaned)
        if _contains_contact_value(cleaned):
            contact_values.append(cleaned)
    return "\n".join(dict.fromkeys([*name_candidates[:1], *contact_values])).strip()


def _contains_contact_value(value: str) -> bool:
    return bool(
        re.search(r"[A-Z0-9._%+\-]+@[A-Z0-9.\-]+\.[A-Z]{2,}", value, flags=re.I)
        or re.search(r"https?://|(?:www\.)?(?:linkedin\.com|github\.com)/", value, flags=re.I)
        or any(_looks_like_phone_number(match.group(0)) for match in re.finditer(r"(?<!\d)\+?[\d][\d\s().-]{7,}\d(?!\d)", value))
    )


def _looks_like_phone_number(value: str) -> bool:
    if re.fullmatch(r"\s*(?:19|20)\d{2}\s*(?:-|–|—|to)\s*(?:present|current|(?:19|20)\d{2})\s*", value, flags=re.I):
        return False
    digits = re.sub(r"\D", "", value)
    return 10 <= len(digits) <= 15


def _looks_like_name(value: str) -> bool:
    if not value.isupper():
        return False
    if not (2 <= len(value.split()) <= 5 and len(value) <= 60):
        return False
    if re.search(r"@|https?://|linkedin|github|\d|\b(?:phone|email|contact|education|skills|projects|experience|summary|objective|certification|achievement)\b", value, flags=re.I):
        return False
    return all(part.replace("'", "").replace("-", "").isalpha() for part in value.split())


def _recover_education(lines: list[str]) -> str:
    terms = re.compile(r"\b(?:b\.?tech|bachelor|master|university|college|school|secondary|cgpa|gpa|marks?|percentage|semester)\b", flags=re.I)
    heading_keys = {_heading_key(alias) for names in SECTION_ALIASES.values() for alias in names}
    recovered: list[str] = []
    education_heading = next((index for index, line in enumerate(lines) if _heading_key(line) == "education"), len(lines))
    candidate_lines = lines[:education_heading + 1]
    for index, line in enumerate(candidate_lines):
        if terms.search(line):
            for nearby in candidate_lines[max(0, index - 1):min(len(candidate_lines), index + 2)]:
                if nearby and nearby not in recovered and _heading_key(nearby) not in heading_keys:
                    recovered.append(nearby)
    return "\n".join(recovered[:18]).strip()


def detect_sections(text: str) -> tuple[str, ...]:
    """Return canonical sections found from heading-like lines, not loose body substrings."""
    found: list[str] = []
    aliases = {
        section: tuple(_heading_key(alias) for alias in names)
        for section, names in SECTION_ALIASES.items()
    }
    for raw_line in text.splitlines():
        line = re.sub(r"^[\s\-–—•▪●*#|]+", "", raw_line).strip()
        if not line:
            continue
        heading_part = re.split(r"\s*[:|]\s*", line, maxsplit=1)[0]
        heading_key = _heading_key(heading_part)
        for section, section_aliases in aliases.items():
            if heading_key in section_aliases:
                if section not in found:
                    found.append(section)
                break
            if len(line) <= 80 and any(heading_key.startswith(f"{alias} ") for alias in section_aliases):
                if section not in found:
                    found.append(section)
                break
    normalized_text = unicodedata.normalize("NFKC", text)
    for section, section_aliases in aliases.items():
        if section in found:
            continue
        for alias in section_aliases:
            pattern = rf"(?:^|[.!?]\s+)({re.escape(alias)})(?=\s*(?:[:|\-–—]|[A-Z0-9]))"
            if re.search(pattern, normalized_text, flags=re.I | re.M):
                found.append(section)
                break
    return tuple(found)


def _heading_key(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", value.casefold()).strip()


def extract_skills(text: str) -> tuple[str, ...]:
    found = [canonicalize_skill(skill) for skill in sorted(KNOWN_SKILLS) if _contains_term(text, skill)]
    sections = extract_sections(text)
    found.extend(_skills_from_explicit_section(sections.get("skills", "")))
    return tuple(dict.fromkeys(found))


def _skills_from_explicit_section(text: str) -> list[str]:
    """Extract market-specific skills from an explicit skills section."""
    found: list[str] = []
    for line in _skill_section_candidate_lines(text):
        if not line:
            continue
        if _is_generic_skill_label(line):
            continue
        if ":" in line:
            label, line = line.split(":", 1)
            if 1 <= len(label.split()) <= 5 and not _is_generic_skill_label(label):
                found.append(canonicalize_skill(label))
        for value in re.split(r"\s*[•▪●|;,]\s*", line):
            cleaned = _clean_skill_candidate(value)
            if _is_skill_like_value(cleaned):
                found.append(canonicalize_skill(cleaned))
    return found[:60]


def _skill_section_candidate_lines(text: str) -> list[str]:
    lines = [
        re.sub(r"^[\s•▪●\uf0b7*\-]+", "", raw_line).strip()
        for raw_line in text.splitlines()
        if raw_line.strip()
    ]
    merged: list[str] = []
    for line in lines:
        cleaned = normalize_ats_text(line)
        key = _skill_key(cleaned)
        has_separator = bool(re.search(r"[:•▪●|;,]", cleaned))
        if (
            merged
            and not has_separator
            and key not in GENERIC_SKILL_LABELS
            and len(cleaned.split()) <= 3
            and not re.search(r"[.!?]$", merged[-1])
        ):
            merged[-1] = f"{merged[-1]} {cleaned}".strip()
            continue
        merged.append(cleaned)
    return merged


def _clean_skill_candidate(value: str) -> str:
    cleaned = normalize_ats_text(value.strip(" .()"))
    cleaned = re.sub(r"\s+", " ", cleaned)
    if cleaned.count("(") > cleaned.count(")"):
        cleaned = re.sub(r"\s*\([^)]*$", "", cleaned).strip()
    return cleaned.strip(" .()")


def _is_skill_like_value(value: str) -> bool:
    if not value or _is_generic_skill_label(value):
        return False
    if len(value) > 70 or len(value.split()) > 7:
        return False
    key = _skill_key(value)
    if key in CANONICAL_SKILL_LABELS or key in {_skill_key(skill) for skill in KNOWN_SKILLS}:
        return True
    if len(value.split()) == 1:
        return bool(re.fullmatch(r"[A-Z0-9]{2,6}", value) or re.fullmatch(r"[A-Z][A-Za-z0-9.+#-]{1,18}", value) and value in {"Bloomberg", "FactSet", "Dealogic", "Euromonitor", "Gartner", "Factiva", "FinBERT", "DSPy", "UiPath", "Copilot", "Perplexity"})
    return bool(re.search(r"\b(?:ai|ml|data|research|analysis|analytics|financial|finance|market|strategy|modelling|engineering|automation|intelligence)\b", value, flags=re.I))


def canonicalize_skill(value: str) -> str:
    cleaned = normalize_ats_text(value.strip(" .()"))
    key = _skill_key(cleaned)
    if key in CANONICAL_SKILL_LABELS:
        return CANONICAL_SKILL_LABELS[key]
    if cleaned.isupper() or re.search(r"[A-Z]{2,}|[+#/.-]", cleaned):
        return cleaned
    return " ".join(part.capitalize() if part.lower() not in {"and", "of", "for"} else part.lower() for part in cleaned.split())


def _contains_term(text: str, term: str) -> bool:
    haystack = " ".join(_skill_tokens(text))
    needle = " ".join(_skill_tokens(term))
    return bool(needle and f" {needle} " in f" {haystack} ")


def _skill_tokens(value: str) -> list[str]:
    normalized = unicodedata.normalize("NFKC", value).casefold()
    normalized = normalized.replace("c++", "c plus plus").replace("ci/cd", "ci cd").replace("pl/sql", "pl sql")
    return re.findall(r"[a-z0-9]+", normalized)


def _skill_key(value: str) -> str:
    return " ".join(_skill_tokens(value))


def _is_generic_skill_label(value: str) -> bool:
    return _skill_key(value) in GENERIC_SKILL_LABELS


def extract_roles(text: str) -> tuple[str, ...]:
    lower = text.lower()
    roles = [role.title() for role in ROLE_HINTS if role in lower]
    return tuple(dict.fromkeys(roles))


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("Install PDF support with: pip install -e '.[docs]'") from exc

    try:
        reader = PdfReader(str(path))
    except Exception as exc:
        raise ValueError("The PDF could not be opened. Re-export it as a standard text-based PDF or upload the DOCX version.") from exc
    pages: list[str] = []
    for page in reader.pages:
        standard_text = ""
        layout_text = ""
        try:
            standard_text = page.extract_text() or ""
        except Exception:
            pass
        try:
            layout_text = page.extract_text(extraction_mode="layout") or ""
        except (TypeError, ValueError, KeyError):
            pass
        pages.append(_select_pdf_extraction(standard_text, layout_text))
    text = "\n".join(pages)
    if not text.strip():
        raise ValueError(
            "No readable text was found in this PDF. It appears to be scanned or image-only. "
            "Run OCR or export/upload the original DOCX so the ATS can analyze the resume accurately."
        )
    return text


def _select_pdf_extraction(standard: str, layout: str) -> str:
    """Prefer logical reading order over visual columns when both are usable."""
    if not standard.strip():
        return layout
    if not layout.strip():
        return standard

    def quality(value: str) -> int:
        heading_keys = {
            _heading_key(alias)
            for names in SECTION_ALIASES.values()
            for alias in names
        }
        lines = [line.strip() for line in value.splitlines() if line.strip()]
        headings = sum(_heading_key(line) in heading_keys for line in lines)
        bullets = value.count("\uf0b7") + value.count("•")
        fragmented = len(re.findall(r"\b(?:[A-Z]\s+){3,}[A-Z]+\b", value))
        excessive_gaps = sum(bool(re.search(r"\S\s{8,}\S", line)) for line in lines)
        return headings * 12 + min(bullets, 20) - fragmented * 8 - excessive_gaps

    # Standard extraction usually follows the PDF content stream and therefore
    # keeps two-column sections intact. Layout mode wins only when it provides a
    # materially clearer structure.
    return layout if quality(layout) > quality(standard) + 4 else standard


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("Install DOCX support with: pip install -e '.[docs]'") from exc

    document = Document(str(path))
    blocks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            blocks.append(" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
    return "\n".join(blocks)
