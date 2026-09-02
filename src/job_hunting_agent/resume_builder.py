from __future__ import annotations

import re
from pathlib import Path

from .models import AtsReport, CandidateProfile, Resume
from .resume import canonicalize_skill, normalize_ats_text

ACTION_VERB_REPLACEMENTS = {
    "building": "Built",
    "conducting": "Conducted",
    "employing": "Employed",
    "utilizing": "Utilized",
    "contributing": "Contributed",
    "spearheading": "Spearheaded",
    "transforming": "Transformed",
    "managing": "Managed",
    "developing": "Developed",
    "designing": "Designed",
    "implementing": "Implemented",
}


def write_improved_resume(resume: Resume, report: AtsReport, profile: CandidateProfile, data_dir: str | Path) -> dict[str, str]:
    output_dir = Path(data_dir) / "improved_resume"
    output_dir.mkdir(parents=True, exist_ok=True)
    detected_name = _candidate_name(resume.text)
    base_name = _safe_filename(profile.name or detected_name or "candidate")
    docx_path = output_dir / f"{base_name}-ATS-Friendly-Resume.docx"

    sections = _resume_sections(resume, report, profile)
    display_name = profile.name.strip() or detected_name or "Candidate"
    _write_docx(docx_path, sections, display_name)
    return {"docx_path": str(docx_path)}


def _resume_sections(resume: Resume, report: AtsReport, profile: CandidateProfile) -> list[tuple[str, list[str]]]:
    target_roles = _ordered_terms((*profile.target_roles, *resume.inferred_roles))
    # Never inject a missing ATS keyword unless it is already supported by the
    # candidate profile or resume evidence.
    skills = _ordered_skills((*profile.skills, *resume.inferred_skills))
    parsed_sections = resume.sections or {}
    experience_source = (
        f"EXPERIENCE\n{parsed_sections['experience']}"
        if parsed_sections.get("experience")
        else (resume.text if _has_experience_heading(resume.text) else "")
    )
    experience_lines = _sanitize_experience_items(_experience_items(experience_source, target_roles))
    experience_bullets = [item.removeprefix("BULLET::") for item in experience_lines if item.startswith("BULLET::")]
    education_lines = _sanitize_section_items(
        _education_section_items(parsed_sections.get("education", "")) or _education_lines(resume.text),
        "EDUCATION",
    )
    project_lines = _sanitize_section_items(_project_section_items(parsed_sections.get("projects", "")), "PROJECTS")
    certifications = _sanitize_section_items(
        _list_section_items(parsed_sections.get("certifications", "")) or _certification_lines(resume.text),
        "CERTIFICATIONS",
    )
    achievements = _sanitize_section_items(
        _list_section_items(parsed_sections.get("achievements", "")) or _achievement_lines(resume.text, experience_bullets),
        "ACHIEVEMENTS",
    )
    languages = _language_lines(resume.text)
    technical_skill_lines = _skill_section_lines(parsed_sections.get("skills", ""), skills)
    original_summary = (resume.sections or {}).get("summary", "").strip()
    summary_items = _summary_section_items(original_summary, resume, profile, target_roles, skills, achievements, experience_bullets)
    contact = _contact_line(profile, parsed_sections.get("contact", ""))

    languages = _language_lines(parsed_sections.get("languages", "")) or languages
    publications = _sanitize_section_items(_section_items(parsed_sections.get("publications", ""), join_wrapped=True), "PUBLICATIONS")
    volunteering = _sanitize_section_items(_section_items(parsed_sections.get("volunteering", ""), join_wrapped=True), "VOLUNTEERING")
    interests = _sanitize_section_items(_section_items(parsed_sections.get("interests", ""), join_wrapped=True), "INTERESTS")
    teaching_vision = _prose_section_items(parsed_sections.get("teaching_vision", ""))
    teaching_subjects = _sanitize_section_items(_list_section_items(parsed_sections.get("teaching_subjects", "")), "SUBJECTS")
    core_terms = _high_signal_terms(_compact_term_items(parsed_sections.get("core_competencies", "")))
    soft_terms = _high_signal_terms(_compact_term_items(parsed_sections.get("soft_skills", "")) or _soft_skill_lines(resume.text), soft=True)
    core_competencies = [", ".join(core_terms)] if len(core_terms) >= 3 else []
    soft_skills = [", ".join(soft_terms)] if len(soft_terms) >= 3 else []

    return [
        ("CONTACT", [contact] if contact else []),
        ("PROFESSIONAL SUMMARY", summary_items),
        ("TEACHING VISION", teaching_vision),
        ("SUBJECTS AVAILABLE TO TEACH", teaching_subjects),
        ("CORE COMPETENCIES", core_competencies),
        ("TECHNICAL SKILLS", technical_skill_lines or skills[:12]),
        ("SOFT SKILLS", soft_skills),
        ("PROFESSIONAL EXPERIENCE", experience_lines),
        ("PROJECTS", project_lines),
        ("EDUCATION", education_lines),
        ("CERTIFICATIONS", certifications),
        ("ACHIEVEMENTS", achievements),
        ("PUBLICATIONS & RESEARCH", publications),
        ("VOLUNTEER & LEADERSHIP EXPERIENCE", volunteering),
        ("LANGUAGES", languages),
        ("INTERESTS", interests),
    ]


def _write_docx(path: Path, sections: list[tuple[str, list[str]]], display_name: str) -> None:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("Install DOCX support with: pip install -e '.[docs]'") from exc

    document = Document()
    _configure_document(document)
    title = document.add_paragraph()
    title.style = document.styles["ResumeName"]
    title.add_run(display_name)
    for heading, items in sections:
        if not items:
            continue
        if heading == "CONTACT":
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            paragraph = document.add_paragraph(items[0], style="ResumeBody")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.style = document.styles["ResumeContact"]
            continue
        heading_paragraph = document.add_paragraph(heading, style="ResumeHeading")
        _add_bottom_border(heading_paragraph)
        item_index = 0
        while item_index < len(items):
            item = items[item_index]
            if heading == "PROFESSIONAL SUMMARY" and len(items) > 1:
                paragraph = document.add_paragraph(_clean_sentence(item), style="ResumeBullet")
                paragraph.paragraph_format.keep_together = True
            elif heading in {"CONTACT", "PROFESSIONAL SUMMARY", "TEACHING VISION", "TECHNICAL SKILLS", "CORE COMPETENCIES", "SOFT SKILLS"}:
                paragraph = document.add_paragraph(style="ResumeBody")
                _add_labelled_text(paragraph, item)
            elif heading == "PROFESSIONAL EXPERIENCE":
                if item.startswith("ROLE::"):
                    paragraph = document.add_paragraph(item.removeprefix("ROLE::"), style="ResumeRole")
                    paragraph.runs[0].bold = True
                elif item.startswith("META::"):
                    meta = item.removeprefix("META::")
                    if item_index + 1 < len(items) and items[item_index + 1].startswith("META::") and _looks_like_date_range(items[item_index + 1]):
                        paragraph = document.add_paragraph(style="ResumeMeta")
                        paragraph.add_run(meta)
                        paragraph.add_run("\t" + items[item_index + 1].removeprefix("META::"))
                        item_index += 1
                    else:
                        document.add_paragraph(meta, style="ResumeMeta")
                elif item.startswith("BULLET::"):
                    bullet = _clean_sentence(item.removeprefix("BULLET::"))
                    if _is_resume_metadata(bullet):
                        paragraph = document.add_paragraph(style="ResumeMeta")
                        _add_labelled_text(paragraph, bullet)
                    else:
                        paragraph = document.add_paragraph(bullet, style="ResumeBullet")
                    paragraph.paragraph_format.keep_together = True
                else:
                    document.add_paragraph(_clean_sentence(item), style="ResumeBody")
            else:
                paragraph = document.add_paragraph(style="ResumeBullet")
                _add_labelled_text(
                    paragraph,
                    _clean_sentence(item),
                    bold_dash_prefix=heading in {"EDUCATION", "PROJECTS", "PUBLICATIONS & RESEARCH"},
                )
                paragraph.paragraph_format.keep_together = True
            item_index += 1
    document.save(path)


def _configure_document(document) -> None:
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Mm, Pt, RGBColor

    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.05

    for name, font_size, bold, before, after, color in (
        ("ResumeName", 20, True, 0, 2, "17365D"),
        ("ResumeContact", 8.8, False, 0, 5, "44546A"),
        ("ResumeHeading", 10.5, True, 7, 2, "17365D"),
        ("ResumeRole", 9.8, True, 2, 0, "1F1F1F"),
        ("ResumeMeta", 9, False, 0, 1.5, "595959"),
        ("ResumeBody", 9.5, False, 0, 2.5, "1F1F1F"),
        ("ResumeBullet", 9.3, False, 0, 1.2, "1F1F1F"),
    ):
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(font_size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.05

    styles["ResumeBullet"].base_style = styles["List Bullet"]
    styles["ResumeBullet"].paragraph_format.left_indent = Inches(0.18)
    styles["ResumeBullet"].paragraph_format.first_line_indent = Inches(-0.12)
    styles["ResumeBullet"]._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    styles["ResumeBullet"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    styles["ResumeName"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles["ResumeContact"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles["ResumeRole"].paragraph_format.keep_with_next = True
    styles["ResumeHeading"].paragraph_format.keep_with_next = True

    usable_width = section.page_width - section.left_margin - section.right_margin
    styles["ResumeMeta"].paragraph_format.tab_stops.add_tab_stop(usable_width, 2)

    document.core_properties.title = "ATS-Friendly Resume"
    document.core_properties.subject = "Professional resume"


def _add_bottom_border(paragraph) -> None:
    """Add the thin section rule used by the supplied Canva reference."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    paragraph_properties = paragraph._p.get_or_add_pPr()
    borders = paragraph_properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        paragraph_properties.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "5")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "4472C4")
    borders.append(bottom)


def _add_labelled_text(paragraph, text: str, *, bold_dash_prefix: bool = False) -> None:
    """Bold compact labels/titles while keeping content as plain ATS-readable text."""
    separator = ":" if ":" in text else (" - " if bold_dash_prefix and " - " in text else "")
    if not separator:
        paragraph.add_run(text)
        return
    prefix, suffix = text.split(separator, 1)
    lead = paragraph.add_run(prefix.strip() + (":" if separator == ":" else ""))
    lead.bold = True
    paragraph.add_run((" - " if separator == " - " else " ") + suffix.strip())


def _summary_section_items(
    original: str,
    resume: Resume,
    profile: CandidateProfile,
    target_roles: list[str],
    skills: list[str],
    achievements: list[str],
    experience_bullets: list[str],
) -> list[str]:
    role = target_roles[0] if target_roles else _role_from_text(resume.text)
    role = role or "Technology professional"
    summary_skills = _summary_skill_terms(skills)
    skill_text = _join_terms(summary_skills)
    experience_phrase = _experience_phrase(resume.text, profile.experience_years)
    focus_terms = _non_overlapping_focus_terms(_domain_focus_terms(resume.text, skills, role, summary_skills), summary_skills)
    measurable = _measurable_summary_evidence([*achievements, *experience_bullets])

    clauses: list[str] = []
    if skill_text:
        clauses.append(f"{role} with {experience_phrase} experience in {skill_text}")
    else:
        clauses.append(f"{role} with {experience_phrase} experience across the documented resume scope")
    if focus_terms:
        clauses.append(f"focused on {', '.join(focus_terms[:4])}")
    if measurable:
        clauses.append(measurable)
    summary = ". ".join(_capitalize_sentence(_clean_sentence(clause).rstrip(".")) for clause in clauses if clause).strip()
    if summary:
        summary += "."

    source_sentence = _best_original_summary_sentence(original, target_roles, skills)
    if source_sentence and len(summary.split()) < 22:
        summary = f"{summary.rstrip('.')} with background in {source_sentence.rstrip('.')}."
    return [_trim_words(summary, 86)]


def _summary_skill_terms(skills: list[str]) -> list[str]:
    strong_markers = re.compile(
        r"\b(?:ai|ml|machine learning|deep learning|data|spark|sql|python|aws|azure|cloud|"
        r"financial|finance|investment|market|research|modelling|strategy|analytics|prompt|automation)\b",
        flags=re.I,
    )
    weak_when_crowded = {"Excel", "Matplotlib", "NumPy", "Pandas", "Jupyter Notebook"}
    strong = [skill for skill in skills if strong_markers.search(skill)]
    selected = strong or skills
    if len(selected) > 4:
        selected = [skill for skill in selected if skill not in weak_when_crowded] or selected
    return selected[:4]


def _best_original_summary_sentence(original: str, target_roles: list[str], skills: list[str]) -> str:
    if not original.strip():
        return ""
    candidates = [
        _clean_sentence(item)
        for item in re.split(r"(?<=[.!?])\s+|[\n•▪●]+", original)
        if 8 <= len(_clean_sentence(item).split()) <= 28
    ]
    signals = [term.casefold() for term in (*target_roles, *skills[:8]) if term]
    generic = re.compile(r"\b(?:hardworking|dynamic|self motivated|challenging environment|organization growth|seeking opportunity)\b", flags=re.I)
    ranked = sorted(
        (item for item in candidates if not generic.search(item)),
        key=lambda item: sum(signal in item.casefold() for signal in signals),
        reverse=True,
    )
    if not ranked or not sum(signal in ranked[0].casefold() for signal in signals):
        return ""
    if re.search(r"\b(?:navigating the intricate landscape|organization growth|challenging environment)\b", ranked[0], flags=re.I):
        return ""
    return ranked[0]


def _role_from_text(text: str) -> str:
    match = re.search(
        r"\b(?:data engineer|machine learning engineer|software engineer|python developer|backend developer|"
        r"full stack developer|data analyst|data scientist|business analyst|consultant|manager|faculty|professor)\b",
        text,
        flags=re.I,
    )
    return match.group(0).title() if match else ""


def _experience_phrase(text: str, configured_years: float) -> str:
    match = re.search(r"\b(\d+(?:\.\d+)?\+?)\s+years?\b", text, flags=re.I)
    if match:
        return f"{match.group(1)} years of"
    if configured_years:
        if 0 < configured_years < 1:
            return f"{round(configured_years * 12):g} months of"
        return f"{configured_years:g} years of"
    return "hands-on"


def _domain_focus_terms(text: str, skills: list[str], role: str = "", summary_skills: list[str] | None = None) -> list[str]:
    lower = f"{text} {' '.join(skills)}".casefold()
    candidates = (
        ("teaching and curriculum delivery", ("adjunct faculty", "faculty", "teaching", "students", "subjects available")),
        ("investment research", ("investment banking", "m&a", "capital issuance", "valuation", "pitchbook")),
        ("financial modelling", ("financial modelling", "financial models", "finance", "valuation")),
        ("market research", ("market research", "market sizing", "competitive intelligence", "secondary research")),
        ("AI automation", ("generative ai", "agentic ai", "prompt engineering", "rpa", "uipath")),
        ("data pipelines", ("pipeline", "etl", "spark", "airflow", "databricks")),
        ("analytics reporting", ("dashboard", "report", "analytics", "power bi", "tableau")),
        ("backend services", ("api", "fastapi", "django", "flask", "microservice")),
        ("cloud deployment", ("aws", "azure", "gcp", "terraform", "docker", "kubernetes")),
        ("machine learning", ("machine learning", "model", "tensorflow", "pytorch", "scikit-learn")),
        ("stakeholder delivery", ("stakeholder", "client", "vendor", "business user")),
    )
    terms = [label for label, markers in candidates if any(_contains_focus_marker(lower, marker) for marker in markers)]
    profile_skill_text = " ".join(summary_skills or skills[:6]).casefold()
    technical_target = bool(re.search(r"\b(?:engineer|developer|scientist|data analyst|ai|ml)\b", role, flags=re.I))
    finance_supported = bool(re.search(r"\b(?:financial|finance|investment|market|m&a|valuation|research)\b", profile_skill_text))
    if technical_target and not finance_supported:
        terms = [term for term in terms if term not in {"investment research", "financial modelling", "market research"}]
    return terms


def _non_overlapping_focus_terms(focus_terms: list[str], summary_skills: list[str]) -> list[str]:
    skill_keys = [skill.casefold() for skill in summary_skills]
    unique: list[str] = []
    seen: set[str] = set()
    for term in focus_terms:
        key = term.casefold()
        if any(key == skill or key in skill or skill in key for skill in skill_keys):
            continue
        if key not in seen:
            seen.add(key)
            unique.append(term)
    return unique


def _contains_focus_marker(text: str, marker: str) -> bool:
    if re.search(r"[^a-z0-9 ]", marker, flags=re.I):
        return marker.casefold() in text
    return bool(re.search(rf"(?<![a-z0-9]){re.escape(marker.casefold())}(?![a-z0-9])", text))


def _measurable_summary_evidence(lines: list[str]) -> str:
    ranked = sorted(lines, key=_summary_metric_rank, reverse=True)
    for line in ranked:
        cleaned = _clean_sentence(line)
        metric = re.search(
            r"\b\d+(?:\.\d+)?%|\b\d+\+|\b\d+(?:,\d{3})*(?:\.\d+)?\s*(?:users|records|hours|days|months|projects|clients|files|gb|tb|sla|member)\b",
            cleaned,
            flags=re.I,
        )
        if metric:
            concise = _trim_words(cleaned, 22).rstrip(".")
            return f"supported by measurable evidence such as {concise}"
    return ""


def _summary_metric_rank(line: str) -> int:
    cleaned = _clean_sentence(line)
    score = 0
    if re.search(r"\d+(?:\.\d+)?%", cleaned):
        score += 4
    if re.search(r"\b\d+(?:,\d{3})*\+?\s*(?:hours|users|records|clients|member|projects)\b", cleaned, flags=re.I):
        score += 3
    if re.search(r"\b(?:improved|reduced|increased|achieved|saved|led|managed|developed)\b", cleaned, flags=re.I):
        score += 2
    if re.search(r"\b\d+\+?\s+years?\b", cleaned, flags=re.I):
        score -= 2
    return score


def _join_terms(terms: list[str]) -> str:
    if not terms:
        return ""
    if len(terms) == 1:
        return terms[0]
    return ", ".join(terms[:-1]) + f", and {terms[-1]}"


def _trim_words(value: str, maximum: int) -> str:
    words = value.split()
    if len(words) <= maximum:
        return value
    return " ".join(words[:maximum]).rstrip(" ,;:") + "."


def _sanitize_experience_items(items: list[str]) -> list[str]:
    sanitized: list[str] = []
    for item in items:
        prefix = ""
        value = item
        for candidate in ("ROLE::", "META::", "BULLET::"):
            if item.startswith(candidate):
                prefix = candidate
                value = item.removeprefix(candidate)
                break
        cleaned = _clean_sentence(value)
        if not cleaned or _is_cross_section_heading(cleaned):
            continue
        if prefix == "BULLET::":
            cleaned = _strengthen_action_verb(cleaned)
            if _is_cross_section_heading(cleaned) or _looks_like_contact_or_address(cleaned):
                continue
        sanitized.append(f"{prefix}{cleaned}" if prefix else cleaned)
    return _dedupe_prefixed_lines(sanitized)


def _sanitize_section_items(items: list[str], heading: str) -> list[str]:
    sanitized: list[str] = []
    for item in items:
        cleaned = _clean_sentence(item)
        if not cleaned or _is_cross_section_heading(cleaned):
            continue
        if heading != "CONTACT" and _looks_like_contact_or_address(cleaned):
            continue
        if heading == "EDUCATION" and _looks_like_work_timeline(cleaned):
            continue
        if heading in {"ACHIEVEMENTS", "CERTIFICATIONS"} and _is_resume_metadata(cleaned):
            continue
        sanitized.append(cleaned)
    if heading == "EDUCATION":
        sanitized = _dedupe_secondary_education_labels(sanitized)
    return _dedupe_lines(sanitized)


def _is_cross_section_heading(value: str) -> bool:
    key = re.sub(r"[^a-z0-9]+", " ", value.casefold()).strip()
    return key in {
        "academic credentials",
        "academic qualifications",
        "achievements",
        "certification",
        "certifications",
        "contact",
        "contact details",
        "core competencies",
        "education",
        "interests",
        "languages",
        "personal details",
        "professional experience",
        "professional summary",
        "profile summary",
        "projects",
        "publications",
        "skills",
        "soft skills",
        "technical skills",
        "volunteer experience",
        "work experience",
    }


def _looks_like_contact_or_address(value: str) -> bool:
    return bool(
        re.search(r"@|https?://|linkedin\.com|github\.com|^\s*address\s*:", value, flags=re.I)
        or re.fullmatch(r"\+?[\d\s().-]{8,}", value)
    )


def _looks_like_work_timeline(value: str) -> bool:
    return bool(
        "|" in value
        and _looks_like_date_range(value)
        and re.search(r"\b(?:engineer|developer|analyst|consultant|associate|officer|manager|lead|intern|pvt|ltd|limited|inc|corp|services|technologies)\b", value, flags=re.I)
    )


def _dedupe_secondary_education_labels(items: list[str]) -> list[str]:
    high_school_items: list[tuple[int, int]] = []
    for index, item in enumerate(items):
        if not re.search(r"\bhigher secondary education\b", item, flags=re.I):
            continue
        year_match = re.search(r"\b(19|20)\d{2}\b", item)
        high_school_items.append((index, int(year_match.group(0)) if year_match else 0))
    if len(high_school_items) <= 1:
        return items
    latest_index = max(high_school_items, key=lambda item: item[1])[0]
    corrected = list(items)
    for index, _ in high_school_items:
        if index != latest_index:
            corrected[index] = re.sub(r"\bHigher Secondary Education\b", "Secondary Education", corrected[index], flags=re.I)
    return corrected


def _high_signal_terms(terms: list[str], *, soft: bool = False) -> list[str]:
    generic = re.compile(
        r"\b(?:hardworking|punctual|honest|positive attitude|quick learner|self motivated|result oriented|"
        r"responsible|flexible|dedicated|sincere|good communication|good knowledge|basic knowledge|"
        r"domain expertise|technical skills|core skills|tools|technology|methodologies|"
        r"strategic planning|leadership|solutions|software development lifecycle|quality assurance|"
        r"quality control|project management|stakeholder engagements?)\b",
        flags=re.I,
    )
    allowed_soft = {
        "communication",
        "leadership",
        "stakeholder management",
        "people management",
        "problem solving",
        "collaboration",
        "mentoring",
        "decision-making",
        "time management",
        "critical thinking",
        "negotiation",
        "planning",
    }
    cleaned: list[str] = []
    for term in terms:
        value = _clean_sentence(term)
        key = value.casefold()
        if not value or len(value) > 70 or generic.search(value):
            continue
        if soft and key not in allowed_soft:
            continue
        cleaned.append(canonicalize_skill(value) if not soft else value.title())
    return _dedupe_lines(cleaned)[:10]


def _experience_items(text: str, target_roles: list[str]) -> list[str]:
    if not text.strip():
        return []
    entries = [entry for entry in _structured_experience_entries(text) if entry["company"] or entry["dates"]]
    if entries:
        items: list[str] = []
        for entry in entries[:12]:
            if entry["title"]:
                items.append(f"ROLE::{entry['title']}")
            company_meta = " | ".join(part for part in (entry["company"], entry["location"]) if part)
            if company_meta:
                items.append(f"META::{company_meta}")
            if entry["dates"]:
                items.append(f"META::{entry['dates']}")
            items.extend(f"BULLET::{bullet}" for bullet in entry["bullets"][:12])
        if any(item.startswith("ROLE::") or item.startswith("META::") for item in items):
            return items

    bullets = _experience_bullets(text)
    if not bullets:
        return []
    role = target_roles[0] if target_roles else "Relevant Professional Experience"
    return [
        f"ROLE::{role}",
        *(f"BULLET::{bullet}" for bullet in bullets),
    ]


def _has_experience_heading(text: str) -> bool:
    return bool(
        re.search(
            r"(?im)^\s*(?:professional\s+experience|corporate\s+experience|work\s+experience|employment\s+history|career\s+history|experience)(?:\s*\([^\n]+\))?\s*[:|\-–—]?\s*$",
            text,
        )
    )


def _structured_experience_entries(text: str) -> list[dict[str, str | list[str]]]:
    lines = _experience_section_lines(text)
    entries: list[dict[str, str | list[str]]] = []
    current: dict[str, str | list[str]] | None = None

    for line in lines:
        lower = line.lower()
        if lower in {"professional experience", "work experience", "employment history", "career history", "experience", "key result areas"}:
            continue
        if lower in {"personal details"}:
            break

        if _merge_wrapped_experience_meta(current, line):
            continue

        timeline_match = _experience_timeline_match(line)
        if not timeline_match:
            timeline_match = _role_first_timeline_match(line)
        if timeline_match:
            current = {
                "title": timeline_match["title"],
                "company": timeline_match["company"],
                "location": timeline_match["location"],
                "dates": timeline_match["dates"],
                "bullets": [],
            }
            entries.append(current)
            continue

        label_match = re.match(r"^(?:job title|designation|role|position)\s*[:\-]\s*(.+)$", line, flags=re.I)
        if label_match:
            current = _ensure_entry(entries, current)
            title = _clean_sentence(label_match.group(1))
            if current["title"]:
                current["bullets"].append(f"Project/Account: {title}")  # type: ignore[union-attr]
            else:
                current["title"] = title
            continue

        label_match = re.match(r"^(?:company|company name|employer|organization|organisation)\s*[:\-]\s*(.+)$", line, flags=re.I)
        if label_match:
            current = _ensure_entry(entries, current)
            current["company"] = _clean_sentence(label_match.group(1))
            continue

        label_match = re.match(r"^(?:location)\s*[:\-]\s*(.+)$", line, flags=re.I)
        if label_match:
            current = _ensure_entry(entries, current)
            current["location"] = _clean_sentence(label_match.group(1))
            continue

        if _looks_like_date_range(line):
            current = _ensure_entry(entries, current)
            date_match = _find_date_range(line)
            current["dates"] = _clean_sentence(date_match.group(0) if date_match else line)
            company_prefix = _clean_sentence(line[:date_match.start()] if date_match else "").strip(" |,(-")
            if company_prefix and not current["company"] and "|" in company_prefix and current["title"]:
                left, right = [_clean_sentence(part) for part in company_prefix.split("|", 1)]
                current["title"] = _clean_sentence(f"{current['title']} {left}")
                current["company"] = right
            elif company_prefix and not current["company"]:
                current["company"] = company_prefix
            elif company_prefix and isinstance(current["bullets"], list):
                current["bullets"].append(company_prefix)
            continue

        if re.match(r"^(?:development tools?|database|key result areas?)\s*:", line, flags=re.I):
            if current and len(_clean_sentence(line)) >= 20:
                current["bullets"].append(_clean_sentence(line))  # type: ignore[union-attr]
            continue

        if line.startswith("BULLET_LINE::"):
            current = _ensure_entry(entries, current)
            bullet = _strengthen_action_verb(line.removeprefix("BULLET_LINE::"))
            if bullet:
                current["bullets"].append(bullet)  # type: ignore[union-attr]
            continue

        if _looks_like_role(line):
            if current and (current["title"] or current["company"] or current["bullets"]):
                current = None
            current = _ensure_entry(entries, current)
            role_parts = [_clean_sentence(part) for part in line.split("|", 1)]
            current["title"] = role_parts[0]
            if len(role_parts) > 1:
                current["company"] = role_parts[1]
            continue

        if current and current["title"] and not current["company"] and "|" in line:
            left, right = [_clean_sentence(part) for part in line.split("|", 1)]
            if right and _looks_like_company_meta(right):
                current["title"] = _clean_sentence(f"{current['title']} {left}")
                current["company"] = right
                continue

        if current and _looks_like_company_meta(line) and not current["company"]:
            company, location = _split_company_meta(line)
            current["company"] = company
            current["location"] = location
            continue

        cleaned = _clean_sentence(line)
        if current and len(cleaned) >= 3:
            _append_experience_bullet(current, _strengthen_action_verb(cleaned))

    return [entry for entry in entries if entry["title"] or entry["company"] or entry["bullets"]]


def _experience_section_lines(text: str) -> list[str]:
    lines = _resume_lines(text)
    heading_indexes = [
        index
        for index, line in enumerate(lines)
        if line.lower() in {"experience", "professional experience", "work experience", "employment history", "career history"}
    ]
    if not heading_indexes:
        return lines
    start = heading_indexes[-1] + 1
    end = len(lines)
    for index in range(start, len(lines)):
        if lines[index].lower() in {"personal details", "education", "certification", "certifications", "achievement", "achievements", "languages"}:
            end = index
            break
    return lines[start:end]


def _experience_timeline_match(line: str) -> dict[str, str] | None:
    parts = [_clean_sentence(part) for part in re.split(r"\s*\|\s*", line) if _clean_sentence(part)]
    if len(parts) < 2 or not _looks_like_date_range(parts[0]):
        return None
    company = parts[1]
    title = parts[2] if len(parts) >= 3 else ""
    location = ""
    company_location_match = re.match(r"(.+?),\s*([^,]+)$", company)
    if company_location_match and not re.search(r"\b(?:ltd|limited|pvt|private|inc|corp|corporation)\b$", company, flags=re.I):
        company = _clean_sentence(company_location_match.group(1))
        location = _clean_sentence(company_location_match.group(2))
    return {
        "dates": parts[0],
        "company": company,
        "location": location,
        "title": title,
    }


def _role_first_timeline_match(line: str) -> dict[str, str] | None:
    """Parse 'Role | Company Jan 2020 - Present' resume timelines."""
    if "|" not in line:
        return None
    date_match = _find_date_range(line)
    if not date_match:
        return None
    prefix = _clean_sentence(line[:date_match.start()]).strip(" |,")
    parts = [_clean_sentence(part) for part in prefix.split("|") if _clean_sentence(part)]
    if len(parts) < 2 or not re.search(
        r"\b(?:officer|engineer|developer|analyst|manager|consultant|architect|lead|director|"
        r"specialist|administrator|scientist|associate|intern|faculty|professor|lecturer|researcher)\b",
        parts[0],
        flags=re.I,
    ):
        return None
    return {
        "title": parts[0],
        "company": " | ".join(parts[1:]),
        "location": "",
        "dates": _clean_sentence(date_match.group(0)),
    }


def _find_date_range(value: str):
    month = r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*"
    return re.search(
        rf"\b(?:{month}\s+)?\d{{4}}\s*(?:-|–|—|to)?\s*(?:present|current|(?:{month}\s+)?\d{{4}})\b",
        value,
        flags=re.I,
    )


def _merge_wrapped_experience_meta(current: dict[str, str | list[str]] | None, line: str) -> bool:
    if not current or current.get("dates"):
        return False
    if not current.get("title") and not current.get("company"):
        return False
    if current.get("title") and not current.get("company") and "|" in line:
        timeline = _role_first_timeline_match(_clean_sentence(f"{current['title']} {line}"))
        if timeline:
            current["title"] = timeline["title"]
            current["company"] = timeline["company"]
            current["location"] = timeline["location"]
            current["dates"] = timeline["dates"]
            return True
    combined = _clean_sentence(" ".join(part for part in (str(current.get("company", "")), line) if part))
    date_match = _find_date_range(combined)
    if date_match and (_looks_like_company_meta(combined) or current.get("company")):
        company = _clean_sentence(combined[:date_match.start()]).strip(" (|,-")
        if company:
            current["company"] = company
        current["dates"] = _clean_sentence(date_match.group(0))
        return True
    if current.get("company") and _looks_like_company_meta(line):
        current["company"] = combined
        return True
    return False


def _ensure_entry(
    entries: list[dict[str, str | list[str]]], current: dict[str, str | list[str]] | None
) -> dict[str, str | list[str]]:
    if current is None:
        current = {"title": "", "company": "", "location": "", "dates": "", "bullets": []}
        entries.append(current)
    return current


def _append_experience_bullet(entry: dict[str, str | list[str]], line: str) -> None:
    bullets = entry["bullets"]
    if not isinstance(bullets, list):
        return
    previous = str(bullets[-1]) if bullets else ""
    previous_is_metadata = bool(re.match(r"^(?:development tools?|database|title|project/account)\s*:", previous, flags=re.I))
    is_continuation = bool(
        bullets
        and not previous_is_metadata
        and (
            line[:1].islower()
            or line.lower().startswith(("and ", "or ", "within ", "requirements", "projects", "banking ", "creation "))
            or not previous.endswith((".", ":", ";"))
        )
    )
    if is_continuation:
        bullets[-1] = _clean_sentence(f"{bullets[-1]} {line}")
    else:
        bullets.append(line)


def _section_text(text: str, headings: tuple[str, ...]) -> str:
    pattern = "|".join(re.escape(heading) for heading in headings)
    match = re.search(
        rf"\b(?:{pattern})\b(?P<body>.*?)(?=\b(?:projects|technical skills|education|certifications|achievements|languages|personal details)\b|$)",
        text,
        flags=re.I | re.S,
    )
    return match.group("body") if match else ""


def _resume_lines(text: str) -> list[str]:
    normalized = text.replace("\uf0b7", "\n• ")
    normalized = re.sub(
        r"\b(EXPERIENCE|PROFESSIONAL EXPERIENCE|WORK EXPERIENCE|EMPLOYMENT HISTORY|CAREER HISTORY|PROJECTS|TECHNICAL SKILLS|EDUCATION|CERTIFICATION|CERTIFICATIONS|ACHIEVEMENT|ACHIEVEMENTS|LANGUAGES|PROFILE SUMMARY)\b",
        r"\n\1\n",
        normalized,
        flags=re.I,
    )
    normalized = re.sub(r"\s*([•])\s*", r"\n\1 ", normalized)
    lines: list[str] = []
    for raw_line in normalized.splitlines():
        cleaned = _clean_sentence(raw_line)
        if not cleaned:
            continue
        if cleaned in {".", ";", ":"} and lines:
            lines[-1] = f"{lines[-1]}{cleaned}"
            continue
        if raw_line.lstrip().startswith("•"):
            cleaned = f"BULLET_LINE::{cleaned}"
        lines.append(cleaned)
    return lines


def _section_items(text: str, join_wrapped: bool = False) -> list[str]:
    lines = [_clean_sentence(line) for line in text.splitlines() if _clean_sentence(line)]
    if not join_wrapped:
        return _dedupe_lines(lines)[:12]
    items: list[str] = []
    for line in lines:
        starts_item = bool(re.search(r"\s[-–—]\s|\b(?:intern|developer|engineer|analyst|scientist)\b.*\|", line, flags=re.I))
        if items and not starts_item and not items[-1].endswith((".", ")")):
            items[-1] = _clean_sentence(f"{items[-1]} {line}")
        elif items and not starts_item and line[:1].islower():
            items[-1] = _clean_sentence(f"{items[-1]} {line}")
        else:
            items.append(line)
    return _dedupe_lines(items)[:12]


def _list_section_items(text: str) -> list[str]:
    """Rebuild bullet sections while joining PDF-wrapped continuation lines."""
    items: list[str] = []
    current = ""
    for raw_line in text.splitlines():
        raw = raw_line.strip()
        if not raw:
            continue
        starts_bullet = bool(re.match(r"^[•▪●\uf0b7*\-]\s*", raw))
        cleaned = _clean_sentence(re.sub(r"^[•▪●\uf0b7*\-]\s*", "", raw))
        if not cleaned:
            continue
        if starts_bullet:
            if current:
                items.append(current)
            current = cleaned
        elif current:
            current = _clean_sentence(f"{current} {cleaned}")
        else:
            current = cleaned
    if current:
        items.append(current)
    return _dedupe_lines([item for item in items if not _is_section_heading(item)])[:16]


def _project_section_items(text: str) -> list[str]:
    """Split project sections even when PDF extraction removed bullet markers."""
    lines = [_clean_sentence(line) for line in text.splitlines() if _clean_sentence(line)]
    if not lines:
        return []
    items: list[str] = []
    current = ""

    def flush() -> None:
        nonlocal current
        if current:
            items.append(current)
            current = ""

    for line in lines:
        starts_project = _looks_like_project_start(line)
        if starts_project:
            flush()
            current = line
        elif current:
            current = _clean_sentence(f"{current} {line}")
        else:
            current = line
    flush()
    if len(items) <= 1:
        return _list_section_items(text)
    return _dedupe_lines([item for item in items if not _is_section_heading(item)])[:12]


def _looks_like_project_start(line: str) -> bool:
    if re.match(r"^[•▪●\uf0b7*\-]\s*", line):
        return True
    if " - " not in line:
        return False
    title, description = line.split(" - ", 1)
    if len(title.split()) > 10 or len(title) > 90:
        return False
    if re.search(r"\b(?:build|built|develop|designed|aims?|analy[sz]es|optimizing|predict|recommend|cluster|classif|framework|system|model|project|using)\b", description, flags=re.I):
        return True
    return bool(re.search(r"\((?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec|q[1-4]|\d{4})", description, flags=re.I))


def _prose_section_items(text: str) -> list[str]:
    """Join PDF line wraps while retaining genuine paragraph boundaries."""
    paragraphs: list[str] = []
    current: list[str] = []
    for raw_line in text.splitlines():
        line = _clean_sentence(raw_line)
        if not line:
            if current:
                paragraphs.append(_clean_sentence(" ".join(current)))
                current = []
            continue
        if current and current[-1].endswith((".", "!", "?")) and len(" ".join(current)) >= 100:
            paragraphs.append(_clean_sentence(" ".join(current)))
            current = [line]
        else:
            current.append(line)
    if current:
        paragraphs.append(_clean_sentence(" ".join(current)))
    return _dedupe_lines(paragraphs)[:8]


def _compact_term_items(text: str) -> list[str]:
    """Keep short competency lists without merging neighboring terms."""
    items = _list_section_items(text)
    if len(items) <= 1:
        lines = [_clean_sentence(line) for line in text.splitlines() if _clean_sentence(line)]
        items = [line for line in lines if not _is_section_heading(line)]
    terms: list[str] = []
    for item in items:
        terms.extend(part.strip() for part in re.split(r"\s*[|;•▪●]\s*", item) if part.strip())
    return _dedupe_lines(terms)[:16]


def _education_section_items(text: str) -> list[str]:
    lines = [_clean_sentence(line) for line in text.splitlines() if _clean_sentence(line)]
    parallel_entries = _parallel_education_entries(lines)
    if parallel_entries:
        return parallel_entries
    entries: list[str] = []
    current = ""
    education_signal = re.compile(
        r"\b(?:pursuing|education|b\.?tech|m\.?tech|bachelor|master|mba|ph\.?d|diploma|degree|"
        r"b\.?sc|m\.?sc|b\.?a|m\.?a|executive programme?|executive program|"
        r"university|college|school|institute|academy|iim|iit|cgpa|gpa|grade|marks?|percentage|"
        r"secondary|higher secondary|senior secondary|matriculation|icse|cbse|isc)\b",
        flags=re.I,
    )

    def flush() -> None:
        nonlocal current
        if current:
            entries.append(_clean_sentence(current))
            current = ""

    for line in lines:
        if _looks_like_work_timeline(line):
            break
        if re.fullmatch(r"t\s*h", line, flags=re.I):
            continue
        if "|" in line and not re.search(r"\b(?:university|college|school|institute|academy|iim|iit|icse|cbse|isc)\b", line, flags=re.I):
            continue
        starts_entry = bool(
            re.match(r"^(?:19|20)\d{2}\s*:", line)
            or re.match(
                r"^(?:pursuing|currently|executive programme?|executive program|mba\b|ph\.?d\b|"
                r"b\.?tech\b|m\.?tech\b|b\.?sc\b|m\.?sc\b|bachelor\b|master\b|diploma\b|"
                r"higher secondary\b|secondary education\b|senior secondary\b|matriculation\b)",
                line,
                flags=re.I,
            )
        )
        has_signal = bool(education_signal.search(line))
        if starts_entry:
            flush()
            current = line
        elif current and (has_signal or len(line) <= 90):
            current = _clean_sentence(f"{current} {line}")
        elif has_signal:
            flush()
            current = line
    flush()
    return _dedupe_secondary_education_labels(_dedupe_lines(entries))[:10]


def _parallel_education_entries(lines: list[str]) -> list[str]:
    credential_indexes = [
        index for index, line in enumerate(lines)
        if _looks_like_education_credential(line)
    ]
    if len(credential_indexes) < 2 or credential_indexes[0] <= 2:
        return []
    dates = [line for line in lines[:credential_indexes[0]] if _is_standalone_education_date(line)]
    metrics = _education_metrics(lines[:credential_indexes[0]])
    if not dates:
        return []
    entries: list[str] = []
    for position, start in enumerate(credential_indexes):
        end = credential_indexes[position + 1] if position + 1 < len(credential_indexes) else len(lines)
        credential = lines[start]
        details = [
            line for line in lines[start + 1:end]
            if not _is_isolated_ordinal(line)
            and not _is_standalone_education_date(line)
            and not re.search(r"^(?:marks?|cgpa|gpa)\b", line, flags=re.I)
        ]
        parts = []
        if position < len(dates):
            parts.append(f"{dates[position]}:")
        parts.append(credential)
        if details:
            parts.append("- " + _clean_sentence(" ".join(details)))
        if position < len(metrics):
            parts.append(f"({metrics[position]})")
        entries.append(_clean_sentence(" ".join(parts)))
    return _dedupe_secondary_education_labels(_dedupe_lines(entries))[:10]


def _looks_like_education_credential(line: str) -> bool:
    return bool(
        re.search(
            r"\b(?:b\.?tech|m\.?tech|bachelor|master|mba|ph\.?d|diploma|b\.?sc|m\.?sc|"
            r"higher secondary|senior secondary|secondary\s*\(?x\)?|class\s*(?:x|xii|10|12))\b",
            line,
            flags=re.I,
        )
    )


def _is_standalone_education_date(line: str) -> bool:
    return bool(
        re.fullmatch(r"(?:19|20)\d{2}\s*(?:-|–|—|to)\s*(?:present|current|(?:19|20)\d{2})", line, flags=re.I)
        or re.fullmatch(r"(?:19|20)\d{2}", line)
    )


def _education_metrics(lines: list[str]) -> list[str]:
    metrics: list[str] = []
    index = 0
    while index < len(lines):
        line = lines[index]
        if re.search(r"\bcgpa\b", line, flags=re.I):
            value = line
            if index + 1 < len(lines) and not _is_standalone_education_date(lines[index + 1]):
                value = _clean_sentence(f"{value} {lines[index + 1]}")
                index += 1
            if index + 1 < len(lines) and _is_isolated_ordinal(lines[index + 1]):
                value = _clean_sentence(f"{value}{lines[index + 1]}")
                index += 1
            metrics.append(_repair_ordinal_metric(value))
        elif re.search(r"\bmarks?\b|\bpercentage\b", line, flags=re.I):
            value = line
            if index + 1 < len(lines) and re.search(r"\d+(?:\.\d+)?\s*%", lines[index + 1]):
                value = _clean_sentence(f"{value} {lines[index + 1]}")
                index += 1
            metrics.append(value)
        index += 1
    return metrics


def _repair_ordinal_metric(value: str) -> str:
    cleaned = _clean_sentence(value)
    cleaned = re.sub(r"\b(\d+)\s+(?:Semester\))\s*(st|nd|rd|th)\b", r"\1\2 Semester)", cleaned, flags=re.I)
    cleaned = re.sub(r"\b(\d+)\s*(st|nd|rd|th)\s+Semester\b", r"\1\2 Semester", cleaned, flags=re.I)
    return cleaned


def _is_isolated_ordinal(line: str) -> bool:
    return bool(re.fullmatch(r"(?:st|nd|rd|th|t\s*h)", line, flags=re.I))


def _soft_skill_lines(text: str) -> list[str]:
    vocabulary = (
        "communication", "communicator", "leadership", "leader", "negotiation", "negotiator",
        "planning", "planner", "people management", "stakeholder management", "decision-making",
        "decision maker", "decision-maker", "problem solving", "problem-solving", "collaboration",
        "teamwork", "adaptability", "mentoring", "time management", "critical thinking",
    )
    found: list[str] = []
    for raw_line in text.splitlines():
        line = _clean_sentence(raw_line)
        key = line.casefold()
        if len(line) <= 60 and any(key == term for term in vocabulary):
            found.append(line)
    return _dedupe_lines(found)[:12]


def _looks_like_role(line: str) -> bool:
    role_pattern = r"\b(officer|engineer|developer|analyst|manager|consultant|architect|lead|director|specialist|administrator|scientist|associate|intern|faculty|professor|lecturer|researcher)\b"
    if "|" in line and re.search(role_pattern, line, flags=re.I):
        return True
    if ":" in line or "," in line or "." in line or len(line.split()) > 8:
        return False
    if re.match(
        r"^(?:actively|serving|played|efficiently|utilized|transformed|managed|developed|designed|contributed|implemented|defined|built|led|created|and)\b",
        line,
        flags=re.I,
    ):
        return False
    return bool(
        re.search(
            role_pattern,
            line,
            flags=re.I,
        )
    ) and len(line) <= 90


def _looks_like_company_meta(line: str) -> bool:
    return ("|" in line and len(line) <= 140) or bool(
        re.search(r"\b(?:pvt|private|limited|ltd|inc|corp|corporation|technologies|solutions|systems|services|bank|consultancy)\b", line, flags=re.I)
    )


def _split_company_meta(line: str) -> tuple[str, str]:
    parts = [_clean_sentence(part) for part in line.split("|", 1)]
    return parts[0], parts[1] if len(parts) > 1 else ""


def _looks_like_date_range(line: str) -> bool:
    return bool(
        _find_date_range(line)
        or re.search(r"\b\d{4}\s*(?:-|–|—|to)\s*(?:present|current|\d{4})\b", line, flags=re.I)
    )


def _experience_bullets(text: str) -> list[str]:
    bullets = [_strengthen_action_verb(item) for item in _split_on_resume_bullets(text)]
    preferred = [
        item
        for item in bullets
        if any(term in item.lower() for term in ("spark", "scala", "python", "sql", "data", "etl", "pipeline", "workflow", "analysis", "report", "architecture"))
    ]
    selected = _dedupe_lines(preferred or bullets)
    return selected[:10]


def _project_bullets(text: str, experience_lines: list[str]) -> list[str]:
    projects = _project_account_lines(text)
    if projects:
        return projects
    project_terms = ("project", "framework", "solution", "deployment", "jenkins", "autosys", "databricks")
    bullets = [
        _strip_section_prefix(_strengthen_action_verb(item))
        for item in _split_on_resume_bullets(_section_text(text, ("projects",)))
        if any(term in item.lower() for term in project_terms)
    ]
    return _dedupe_lines([item for item in bullets if item not in experience_lines and not _is_resume_metadata(item)])[:5]


def _project_account_lines(text: str) -> list[str]:
    accounts = []
    for line in _experience_section_lines(text):
        match = re.match(r"^title\s*:\s*(.+)$", line, flags=re.I)
        if not match:
            continue
        value = _clean_sentence(match.group(1))
        if value and value.lower() not in {"gft", "corporate technology solution"}:
            accounts.append(f"{value}: delivered technology solutions, data workflows, and stakeholder-facing outcomes in a client project/account environment.")
    return _dedupe_lines(accounts)[:6]


def _education_lines(text: str) -> list[str]:
    matches = []
    patterns = (
        r"Executive Programme on Business Analytics using AI/ML from IIM Calcutta",
        r"B\.?Tech\.? in Computer Science from .*?University of Technology",
        r"Higher Secondary Education from St\.? Jude'?s High School,? ICSE Board",
    )
    for pattern in patterns:
        for match in re.finditer(pattern, text, flags=re.I):
            matches.append(_clean_sentence(match.group(0)))
    return _dedupe_lines(matches)[:4]


def _certification_lines(text: str) -> list[str]:
    compact = _clean_sentence(text)
    lines = []
    patterns = (
        r"NSE.?s Certification on Financial Market\s*\(Basic Module\)",
        r"NSE.?s Certification on Securities Market",
        r"NSE.?s Certification on Mutual Funds",
        r"Professional Training from Petaa Bytes Institute on Hadoop, PIG, Hive & Sqoop, Spark with Scala",
        r"Completed Databricks certification on Apache Spark\s*\(ETL Extraction Series, Cluster Setup on AWS\)",
    )
    for pattern in patterns:
        for match in re.finditer(pattern, compact, flags=re.I):
            lines.append(_clean_sentence(match.group(0)))
    return _dedupe_lines(lines)[:5]


def _achievement_lines(text: str, experience_lines: list[str]) -> list[str]:
    recognitions = _recognition_lines(text)
    if recognitions:
        return recognitions
    measurable_terms = (
        "%",
        "years",
        "csv",
        "json",
        "parquet",
        "reduced",
        "optimized",
        "improved",
        "automated",
        "streamlined",
        "enhanced",
        "delivered",
        "implemented",
    )
    achievements = [
        line
        for line in experience_lines
        if not _is_resume_metadata(line)
        and (re.search(r"\d", line) or any(term in line.lower() for term in measurable_terms))
    ]
    if not achievements and experience_lines:
        achievements = [line for line in experience_lines if not _is_resume_metadata(line)][:2]
    return _dedupe_lines(achievements)[:4]


def _recognition_lines(text: str) -> list[str]:
    lines = []
    compact = _clean_sentence(text)
    batch_topper = re.search(r"completed Executive Programme on Business Analytics at IIM Calcutta as Batch Topper", compact, flags=re.I)
    if batch_topper:
        lines.append("Completed Executive Programme on Business Analytics at IIM Calcutta as Batch Topper.")
    award_match = re.search(
        r"recognized for outstanding performance with (?P<body>.*?Star Employee Award from Capgemini India Ltd in 2020)",
        compact,
        flags=re.I,
    )
    if award_match:
        body = _clean_sentence(award_match.group("body"))
        body = re.sub(r"^various awards and appreciations including\s+", "", body, flags=re.I)
        for item in re.split(r",\s+and\s+|,\s*(?=Pat on Back|Star Employee|Best Debutant)", body):
            cleaned = _clean_sentence(item)
            if cleaned and re.search(r"\baward|appreciation|employee|debutant|pat on back\b", cleaned, flags=re.I):
                lines.append(cleaned)
    return _dedupe_lines(lines)[:5]


def _language_lines(text: str) -> list[str]:
    match = re.search(r"\blanguages known\s*[:\-]\s*([^•\n]{3,160})", text, flags=re.I)
    if not match:
        match = re.search(r"(?im)^languages\s*[:\-]\s*([^•\n]{3,160})$", text)
    if match:
        language_text = match.group(1)
    else:
        nonempty = [line for line in text.splitlines() if line.strip()]
        if not nonempty or len(nonempty) > 3:
            return []
        language_text = nonempty[0]
    language_text = re.split(r"\b(?:address|education|experience|skills|projects)\b", language_text, maxsplit=1, flags=re.I)[0]
    language_text = re.sub(r"\s+\band\b\s+", ",", language_text, flags=re.I)
    values = [item.strip(" .;") for item in re.split(r"[,/|]", language_text) if item.strip(" .;")]
    return _dedupe_lines([_clean_sentence(value) for value in values])[:6]


def _is_resume_metadata(value: str) -> bool:
    return bool(
        re.match(
            r"^(?:database|development tools?|devops tool|cloud technology|big data technology|data visualization tool|operating systems|scripting language|source versioning control tool|programming language|eagle technology|title|project/account)\s*:",
            value,
            flags=re.I,
        )
    )


def _technical_skill_lines(skills: list[str]) -> list[str]:
    categories = {
        "Programming Languages": ("python", "java", "scala", "sql", "javascript", "typescript", "c++", "c#"),
        "Databases": ("postgresql", "postgres", "mysql", "oracle", "sql server", "mongodb", "snowflake", "redshift", "dynamodb"),
        "Cloud": ("aws", "azure", "gcp", "google cloud"),
        "Big Data": ("spark", "hadoop", "hive", "databricks", "airflow", "kafka", "sqoop", "pyspark"),
        "Tools": ("fastapi", "autosys", "jenkins", "git", "github", "docker", "kubernetes", "tableau", "power bi", "excel", "jira"),
        "Analytics": ("machine learning", "data analytics", "risk modeller", "risk browser", "underwriting iq", "matplotlib", "seaborn"),
        "Operating Systems": ("linux", "unix", "windows"),
    }
    remaining = list(skills)
    lines: list[str] = []
    for label, keywords in categories.items():
        matched = []
        for skill in list(remaining):
            key = skill.lower()
            if any(keyword == key or (len(keyword) > 3 and keyword in key) for keyword in keywords):
                matched.append(skill)
                remaining.remove(skill)
        if matched:
            lines.append(f"{label}: {', '.join(_ordered_skills(tuple(matched)))}")
    if remaining:
        lines.append(f"Additional Skills: {', '.join(_ordered_skills(tuple(remaining[:10])))}")
    return lines[:7]


def _skill_section_lines(original: str, inferred: list[str]) -> list[str]:
    """Preserve uploaded skill labels and keywords, then add only supported omissions."""
    if not original.strip():
        return _technical_skill_lines(inferred)
    original_lines = _categorized_skill_lines(original) or _skill_raw_lines(original)
    preserved: list[str] = []
    for line in original_lines:
        if _is_section_heading(line):
            continue
        # Normalize common separators without flattening category labels such as
        # "Cloud: AWS, Azure". Keeping those labels improves scanning by people
        # and prevents niche tools from being dropped by the known-skill list.
        cleaned = re.sub(r"\s*[|;•▪●]\s*", ", ", line)
        cleaned = re.sub(r",\s*,+", ", ", cleaned).strip(" ,")
        if cleaned:
            preserved.append(_normalize_skill_line(cleaned))

    searchable = " ".join(preserved).casefold()
    additions = [skill for skill in inferred if canonicalize_skill(skill).casefold() not in searchable]
    if additions:
        preserved.append(f"Additional Skills: {', '.join(_ordered_skills(tuple(additions)))}")
    return _dedupe_lines(preserved)[:14] if preserved else _technical_skill_lines(inferred)


def _skill_raw_lines(text: str) -> list[str]:
    return _dedupe_lines(
        [
            _clean_sentence(re.sub(r"^[\s•▪●\uf0b7*\-]+", "", line))
            for line in text.splitlines()
            if _clean_sentence(re.sub(r"^[\s•▪●\uf0b7*\-]+", "", line))
        ]
    )


def _normalize_skill_line(line: str) -> str:
    cleaned = normalize_ats_text(_clean_sentence(line))
    if ":" in cleaned:
        label, values = cleaned.split(":", 1)
        terms = _split_skill_terms(values)
        if terms:
            return f"{_clean_skill_label(label)}: {', '.join(_ordered_skills(tuple(terms)))}"
        return f"{_clean_skill_label(label)}: {values.strip()}"
    terms = _split_skill_terms(cleaned)
    return ", ".join(_ordered_skills(tuple(terms))) if terms else cleaned


def _split_skill_terms(value: str) -> list[str]:
    return [
        _clean_skill_term(part)
        for part in re.split(r"\s*(?:,|\||;|•|▪|●)\s*", value)
        if _clean_skill_term(part)
    ]


def _clean_skill_term(value: str) -> str:
    cleaned = normalize_ats_text(_clean_sentence(value)).strip(" .()")
    if cleaned.count("(") > cleaned.count(")"):
        cleaned = re.sub(r"\s*\([^)]*$", "", cleaned).strip()
    cleaned = re.sub(r"\bM&a\b", "M&A", cleaned)
    return cleaned.strip(" .()")


def _clean_skill_label(value: str) -> str:
    label = normalize_ats_text(_clean_sentence(value)).strip(":")
    replacements = {
        "ai and technology": "AI & Technology",
        "ai technology": "AI & Technology",
        "cs fundamentals": "CS Fundamentals",
        "cloud and devops": "Cloud & DevOps",
        "cloud devops": "Cloud & DevOps",
        "data analytics and visualization": "Data Analytics & Visualization",
        "data science and ai ml": "Data Science & AI/ML",
        "devops": "DevOps",
        "domain expertise": "Domain Expertise",
        "big data": "Big Data",
        "data platforms": "Data Platforms",
        "image processing": "Image Processing",
        "mlops and cloud": "MLOps & Cloud",
        "natural language processing": "Natural Language Processing",
        "programming": "Programming",
        "programming languages": "Programming Languages",
        "databases": "Databases",
        "database": "Databases",
        "operating systems": "Operating Systems",
        "research databases": "Research Databases",
    }
    return replacements.get(label.casefold(), label.title() if label.islower() else label)


def _categorized_skill_lines(text: str) -> list[str]:
    lines = [_clean_sentence(line) for line in text.splitlines() if _clean_sentence(line)]
    categories: list[tuple[str, list[str]]] = []
    current_label = ""
    current_values: list[str] = []

    def flush() -> None:
        nonlocal current_label, current_values
        if current_label and current_values:
            categories.append((current_label, current_values.copy()))
        current_label, current_values = "", []

    for line in lines:
        if ":" in line:
            label, values = line.split(":", 1)
            if _looks_like_skill_label(label):
                flush()
                current_label = label
                if values.strip():
                    current_values.append(values.strip())
                continue
        label_candidate = line.rstrip(":")
        is_label = bool(
            _looks_like_skill_label(label_candidate)
            and not re.search(r"[•▪●\uf0b7,;]", label_candidate)
        )
        if is_label:
            flush()
            current_label = label_candidate
        elif current_label:
            current_values.append(line)
    flush()
    return [f"{label}: {_join_skill_values(values)}" for label, values in categories]


def _join_skill_values(values: list[str]) -> str:
    current = ""
    for value in values:
        cleaned = _clean_sentence(value)
        if not cleaned:
            continue
        if not current:
            current = cleaned
        elif _is_wrapped_skill_continuation(current, cleaned):
            current = _clean_sentence(f"{current} {cleaned}")
        else:
            current = _clean_sentence(f"{current} • {cleaned}")
    return current


def _is_wrapped_skill_continuation(current: str, next_value: str) -> bool:
    last_word = current.split()[-1].strip(" ,;:.()").casefold() if current.split() else ""
    return bool(
        current.endswith((",", "/", "&", "("))
        or next_value[:1].islower()
        or (len(next_value.split()) <= 2 and not re.search(r"[•,;:]", next_value))
        or last_word in {"competitive", "operating", "prompt", "financial", "market", "secondary", "stock"}
    )


def _looks_like_skill_label(value: str) -> bool:
    label = _clean_sentence(value)
    return bool(
        len(label) <= 60
        and len(label.split()) <= 7
        and re.search(
            r"\b(?:skills?|tools?|expertise|technology|technologies|databases?|platforms?|"
            r"languages?|frameworks?|systems?|methods?|competencies|devops|cloud|mlops|"
            r"analytics|visualization|processing|fundamentals|domain|science|ai|ml|programming)\b",
            label,
            flags=re.I,
        )
    )


def _is_section_heading(value: str) -> bool:
    key = re.sub(r"[^a-z0-9]+", " ", value.casefold()).strip()
    return key in {
        "skills", "technical skills", "core skills", "key skills",
        "technical competencies", "technologies", "tools and technologies",
    }


def _split_on_resume_bullets(text: str) -> list[str]:
    normalized = text.replace("\uf0b7", " • ")
    parts = re.split(r"\s*[•]\s*", normalized)
    candidates = []
    for part in parts:
        cleaned = _clean_sentence(part)
        if 45 <= len(cleaned) <= 340 and not cleaned.lower().startswith(("education", "personal details", "date of birth")):
            candidates.append(cleaned)
    if len(candidates) >= 3:
        return candidates
    sentences = re.split(r"(?<=[.!?])\s+", normalized)
    return [_clean_sentence(sentence) for sentence in sentences if len(_clean_sentence(sentence)) >= 45]


def _strip_section_prefix(value: str) -> str:
    return re.sub(r"^(?:projects?|professional experience|technical skills)\s+", "", value, flags=re.I).strip()


def _strengthen_action_verb(value: str) -> str:
    cleaned = _clean_sentence(value)
    gerund_actions = {
        "automating": "Automated",
        "building": "Built",
        "creating": "Created",
        "delivering": "Delivered",
        "designing": "Designed",
        "developing": "Developed",
        "implementing": "Implemented",
        "improving": "Improved",
        "managing": "Managed",
        "optimizing": "Optimized",
        "reducing": "Reduced",
    }
    responsible = re.match(r"^responsible for\s+(\w+)\s+(.+)$", cleaned, flags=re.I)
    if responsible:
        replacement = gerund_actions.get(responsible.group(1).casefold())
        if replacement:
            return f"{replacement} {responsible.group(2)}"
    phrase_replacements = (
        (r"^actively participating in\s+", "Contributed to "),
        (r"^serving as a lead for\s+", "Led "),
        (r"^played a pivotal role in the implementation of\s+", "Implemented "),
        (r"^played a pivotal role in\s+", "Implemented "),
        (r"^efficiently performed\s+", "Performed "),
        (r"^preparation of\s+", "Prepared "),
        (r"^worked with\s+", "Applied "),
        (r"^worked on\s+", "Contributed to "),
    )
    for pattern, replacement_text in phrase_replacements:
        if re.search(pattern, cleaned, flags=re.I):
            return re.sub(pattern, replacement_text, cleaned, count=1, flags=re.I)
    cleaned = re.sub(r"^(?:responsible for|worked on)\s+", "Managed ", cleaned, flags=re.I)
    cleaned = re.sub(r"^helped with\s+", "Supported ", cleaned, flags=re.I)
    cleaned = re.sub(r"^involved in\s+", "Contributed to ", cleaned, flags=re.I)
    first = cleaned.split(" ", 1)[0].lower()
    replacement = ACTION_VERB_REPLACEMENTS.get(first)
    if replacement and " " in cleaned:
        return f"{replacement} {cleaned.split(' ', 1)[1]}"
    return cleaned


def _experience_years(text: str) -> str:
    match = re.search(r"(\d+\+?\s+years)", text, flags=re.I)
    return f"{match.group(1)} of " if match else ""


def _contact_line(profile: CandidateProfile, original_contact: str = "") -> str:
    values = [profile.email, profile.phone, profile.linkedin_profile_url, profile.locations[0] if profile.locations else ""]
    values.extend(_extract_contact_values(original_contact))
    seen: set[str] = set()
    deduped: list[str] = []
    for value in values:
        cleaned = value.strip()
        if not cleaned:
            continue
        if "@" in cleaned:
            key = f"email:{cleaned.casefold()}"
        elif re.fullmatch(r"\+?[\d\s().-]{8,}", cleaned):
            digits = re.sub(r"\D", "", cleaned)
            key = f"phone:{digits[-10:]}"
            cleaned = _format_phone(cleaned)
        else:
            key = f"url:{cleaned.casefold().rstrip('/')}"
        if key not in seen:
            seen.add(key)
            deduped.append(cleaned)
    return " | ".join(deduped)


def _extract_contact_values(text: str) -> list[str]:
    values: list[str] = []
    values.extend(re.findall(r"[A-Z0-9._%+\-]+@[A-Z0-9.\-]+\.[A-Z]{2,}", text, flags=re.I))
    values.extend(re.findall(r"https?://[^\s|•]+|(?:www\.)?(?:linkedin\.com|github\.com)/[^\s|•]+", text, flags=re.I))
    values.extend(
        match.group(0).strip()
        for match in re.finditer(r"(?<!\d)\+?[\d][\d\s().-]{7,}\d(?!\d)", text)
        if _looks_like_phone_number(match.group(0))
    )
    return values


def _looks_like_phone_number(value: str) -> bool:
    if _looks_like_date_range(value):
        return False
    digits = re.sub(r"\D", "", value)
    return 10 <= len(digits) <= 15


def _format_phone(value: str) -> str:
    digits = re.sub(r"\D", "", value)
    if len(digits) == 12 and digits.startswith("91"):
        return f"+91 {digits[2:7]} {digits[7:]}"
    return value.strip()


def _candidate_name(contact: str) -> str:
    lines = [_clean_sentence(line) for line in contact.splitlines() if _clean_sentence(line)]
    for require_upper, candidates in ((True, lines), (False, lines[:8])):
        for line in candidates:
            key = re.sub(r"[^a-z0-9]+", " ", line.casefold()).strip()
            if key in {
                "core competencies", "profile summary", "professional summary", "soft skills",
                "technical skills", "career timeline", "work experience", "professional experience",
                "personal details", "contact details", "areas of expertise", "key competencies",
            }:
                continue
            if require_upper and line != line.upper():
                continue
            if (
                2 <= len(line.split()) <= 5
                and len(line) <= 60
                and not re.search(r"@|https?://|linkedin|github|\d|\b(?:phone|email|contact)\b", line, flags=re.I)
                and all(part.replace("'", "").replace("-", "").isalpha() for part in line.split())
            ):
                return line.title() if line.isupper() else line
    return ""


def _clean_sentence(value: str) -> str:
    cleaned = value.replace("\uf0b7", " ").replace("–", "-").replace("—", "-")
    cleaned = re.sub(r"\s+", " ", cleaned)
    cleaned = re.sub(r"\s+([,.;:])", r"\1", cleaned)
    cleaned = _repair_extraction_spacing(cleaned)
    cleaned = normalize_ats_text(cleaned)
    cleaned = _format_metric_numbers(cleaned)
    cleaned = cleaned.strip(" •\t")
    return re.sub(r"^\s*[-]\s*", "", cleaned)


def _capitalize_sentence(value: str) -> str:
    value = value.strip()
    return value[:1].upper() + value[1:] if value else value


def _format_metric_numbers(value: str) -> str:
    def replace(match) -> str:
        return f"{int(match.group(1)):,}"

    return re.sub(
        r"\b(\d{4,})(?=\s*(?:records|users|transactions|files|rows|clients|reports|dashboards|pipelines)\b)",
        replace,
        value,
        flags=re.I,
    )


def _repair_extraction_spacing(value: str) -> str:
    replacements = {
        "inc luding": "including",
        "decision s": "decisions",
        "timelin es": "timelines",
        "standard s": "standards",
        "Datafram es": "Dataframes",
        "v alidate": "validate",
        "vario us": "various",
        "ext racted": "extracted",
        "P arganas": "Parganas",
    }
    cleaned = value
    for source, target in replacements.items():
        cleaned = re.sub(re.escape(source), target, cleaned, flags=re.I)
    return cleaned


def _dedupe_lines(lines: list[str]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for line in lines:
        key = line.lower()
        if key not in seen:
            seen.add(key)
            result.append(line)
    return result


def _dedupe_prefixed_lines(lines: list[str]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for line in lines:
        key = re.sub(r"^(?:ROLE::|META::|BULLET::)", "", line).casefold()
        if key not in seen:
            seen.add(key)
            result.append(line)
    return result


def _ordered_terms(terms: tuple[str, ...]) -> list[str]:
    seen: set[str] = set()
    ordered: list[str] = []
    for term in terms:
        cleaned = term.strip()
        key = cleaned.lower()
        if cleaned and key not in seen:
            seen.add(key)
            ordered.append(cleaned)
    return ordered


def _ordered_skills(terms: tuple[str, ...]) -> list[str]:
    seen: set[str] = set()
    ordered: list[str] = []
    for term in terms:
        cleaned = _clean_skill_term(canonicalize_skill(term))
        key = cleaned.casefold()
        if cleaned and key not in seen and not _is_noisy_skill_term(cleaned):
            seen.add(key)
            ordered.append(cleaned)
    return ordered


def _is_noisy_skill_term(value: str) -> bool:
    if len(value) > 70 or len(value.split()) > 7:
        return True
    if value.count("(") != value.count(")"):
        return True
    key = re.sub(r"[^a-z0-9]+", " ", value.casefold()).strip()
    return key in {
        "additional skills", "core skills", "domain expertise", "technical skills",
        "ai and technology", "research databases", "tools", "technologies",
    }


def _safe_filename(value: str) -> str:
    cleaned = "".join(char if char.isalnum() or char in "-_" else "-" for char in value.strip())
    return cleaned.strip("-")[:80] or "candidate"
